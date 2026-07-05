from __future__ import annotations

import asyncio
import csv
import json
import os
import smtplib
import ssl
import zipfile
from datetime import UTC, datetime
from email.message import EmailMessage
from email.utils import formataddr, getaddresses
from io import StringIO
from pathlib import Path
from typing import Any, Awaitable, Callable
from urllib.parse import urlparse

import httpx
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter

from .schemas import ToolDefinition


class ApprovalRequired(RuntimeError):
    def __init__(self, tool_id: str, arguments: dict[str, Any], reason: str) -> None:
        super().__init__(reason)
        self.tool_id = tool_id
        self.arguments = arguments
        self.reason = reason


ToolHandler = Callable[[dict[str, Any]], Awaitable[Any]]


class ToolRegistry:
    def __init__(
        self,
        approved_folders: list[Path],
        approved_domains: list[str] | None = None,
        python_runtime: Any | None = None,
        mcp_manager: Any | None = None,
    ) -> None:
        self.approved_folders = [path.resolve() for path in approved_folders]
        self.approved_domains = [item.casefold().strip() for item in (approved_domains or []) if item.strip()]
        self.email_config: dict[str, Any] = self._environment_email_config()
        self.python_runtime = python_runtime
        self.mcp_manager = mcp_manager
        for folder in self.approved_folders:
            folder.mkdir(parents=True, exist_ok=True)
        self.definitions = {
            item.id: item for item in [
                ToolDefinition(
                    id="read_file", name="Read file or document",
                    description="Read text, Markdown, JSON, CSV, Word (.docx), or Excel (.xlsx) from the approved workspace.",
                    approval_policy="never", input_schema={"path": "string"},
                ),
                ToolDefinition(
                    id="write_file", name="Write text file",
                    description="Create or replace a text file inside the approved workspace.",
                    approval_policy="always", input_schema={"path": "string", "content": "string"},
                ),
                ToolDefinition(
                    id="create_word", name="Create Word document",
                    description="Create a .docx document with a title and formatted paragraphs in the approved workspace.",
                    approval_policy="always",
                    input_schema={"path": "string", "title": "string", "content": "string"},
                ),
                ToolDefinition(
                    id="create_excel", name="Create Excel workbook",
                    description="Create a safe .xlsx workbook from JSON rows, arrays, or CSV data in the approved workspace.",
                    approval_policy="always",
                    input_schema={"path": "string", "sheet_name": "string", "data": "any"},
                ),
                ToolDefinition(
                    id="search_files", name="Search local documents",
                    description="Search text, Word, and Excel files in the approved workspace.",
                    approval_policy="never", input_schema={"query": "string"},
                ),
                ToolDefinition(
                    id="http_request", name="HTTP request", description="Call an allowlisted HTTPS endpoint.",
                    approval_policy="mutating", input_schema={"url": "string", "method": "string", "body": "object"},
                ),
                ToolDefinition(
                    id="send_email", name="Send email",
                    description="Send from the encrypted SMTP account configured in Settings to the chosen recipient.",
                    approval_policy="always", input_schema={"to": "string", "subject": "string", "body": "string"},
                ),
                ToolDefinition(
                    id="transform", name="Transform data", description="Format JSON, CSV, text, or UTC dates.",
                    approval_policy="never", input_schema={"operation": "string", "value": "any"},
                ),
                ToolDefinition(
                    id="python_code", name="Python code",
                    description="Developer Mode: run a Python 3 snippet with JSON input, a hard timeout, and visible approval.",
                    approval_policy="always",
                    input_schema={"code": "string", "input": "any", "timeout_seconds": "string"},
                ),
                ToolDefinition(
                    id="mcp_call", name="MCP server tool",
                    description="Call one tool on an explicitly enabled local stdio MCP server.",
                    approval_policy="always",
                    input_schema={"server_id": "string", "tool_name": "string", "arguments": "object"},
                ),
            ]
        }
        self.handlers: dict[str, ToolHandler] = {
            "read_file": self._read_file,
            "write_file": self._write_file,
            "create_word": self._create_word,
            "create_excel": self._create_excel,
            "search_files": self._search_files,
            "http_request": self._http_request,
            "send_email": self._send_email,
            "transform": self._transform,
            "python_code": self._python_code,
            "mcp_call": self._mcp_call,
        }

    @staticmethod
    def _environment_email_config() -> dict[str, Any]:
        host = os.environ.get("STUDIO_SMTP_HOST", "").strip()
        username = os.environ.get("STUDIO_SMTP_USERNAME", "").strip()
        password = os.environ.get("STUDIO_SMTP_PASSWORD", "")
        sender = os.environ.get("STUDIO_SMTP_FROM", username).strip()
        if not all((host, username, password, sender)):
            return {}
        port = int(os.environ.get("STUDIO_SMTP_PORT", "465"))
        return {
            "provider": "custom", "host": host, "port": port,
            "security": "ssl" if port == 465 else "starttls",
            "username": username, "password": password, "sender_email": sender, "sender_name": "",
        }

    def configure_email(self, config: dict[str, Any] | None) -> None:
        self.email_config = dict(config or {})

    def email_status(self) -> dict[str, Any]:
        config = self.email_config
        return {
            "configured": bool(config.get("host") and config.get("username") and config.get("password")),
            "provider": config.get("provider", "gmail"),
            "sender_email": config.get("sender_email", ""),
            "sender_name": config.get("sender_name", ""),
            "username": config.get("username", ""),
            "host": config.get("host", ""),
            "port": int(config.get("port", 465)),
            "security": config.get("security", "ssl"),
            "password_saved": bool(config.get("password")),
        }

    def catalog(self) -> list[ToolDefinition]:
        return list(self.definitions.values())

    def requires_approval(self, tool_id: str, arguments: dict[str, Any]) -> bool:
        definition = self.definitions[tool_id]
        if definition.approval_policy == "always":
            return True
        if definition.approval_policy == "mutating":
            return str(arguments.get("method", "GET")).upper() not in {"GET", "HEAD", "OPTIONS"}
        return False

    async def execute(
        self,
        tool_id: str,
        arguments: dict[str, Any],
        *,
        approved: bool = False,
        allow_without_approval: bool = False,
    ) -> Any:
        if tool_id not in self.handlers:
            raise ValueError(f"Unknown tool: {tool_id}")
        automatic_email = tool_id == "send_email" and allow_without_approval
        if self.requires_approval(tool_id, arguments) and not approved and not automatic_email:
            raise ApprovalRequired(tool_id, arguments, f"{self.definitions[tool_id].name} needs your approval")
        return await self.handlers[tool_id](arguments)

    async def _python_code(self, arguments: dict[str, Any]) -> dict[str, Any]:
        if not self.python_runtime:
            raise RuntimeError("Python support is unavailable in this build")
        return await asyncio.to_thread(
            self.python_runtime.execute,
            str(arguments.get("code", "")),
            arguments.get("input"),
            int(arguments.get("timeout_seconds", 30) or 30),
        )

    async def _mcp_call(self, arguments: dict[str, Any]) -> dict[str, Any]:
        if not self.mcp_manager:
            raise RuntimeError("MCP support is unavailable in this build")
        values = arguments.get("arguments", {})
        if not isinstance(values, dict):
            raise ValueError("MCP arguments must be a JSON object")
        try:
            return await asyncio.wait_for(
                self.mcp_manager.call(
                    str(arguments.get("server_id", "")), str(arguments.get("tool_name", "")), values
                ),
                timeout=60,
            )
        except TimeoutError as error:
            raise RuntimeError("The MCP tool exceeded the 60-second limit") from error

    def _safe_path(self, raw: str) -> Path:
        if not raw.strip():
            raise ValueError("Choose a file path inside the approved workspace")
        source = Path(raw).expanduser()
        candidate = (self.approved_folders[0] / source).resolve() if not source.is_absolute() else source.resolve()
        if not any(candidate == root or root in candidate.parents for root in self.approved_folders):
            raise PermissionError("The path is outside the approved workspace folders")
        return candidate

    @staticmethod
    def _bounded_file(path: Path, maximum: int = 10 * 1024 * 1024) -> None:
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"File not found: {path.name}")
        if path.stat().st_size > maximum:
            raise ValueError(f"{path.name} is too large for one workflow step")

    async def _read_file(self, arguments: dict[str, Any]) -> str:
        path = self._safe_path(str(arguments["path"]))
        return await asyncio.to_thread(self._read_file_sync, path)

    def _read_file_sync(self, path: Path) -> str:
        self._bounded_file(path)
        suffix = path.suffix.casefold()
        if suffix == ".docx":
            self._safe_office_archive(path)
            document = Document(path)
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
            return "\n".join(parts)[:200_000]
        if suffix == ".xlsx":
            self._safe_office_archive(path)
            workbook = load_workbook(path, read_only=True, data_only=True)
            payload: dict[str, list[list[Any]]] = {}
            try:
                for sheet in workbook.worksheets[:10]:
                    rows: list[list[Any]] = []
                    for index, row in enumerate(sheet.iter_rows(values_only=True)):
                        if index >= 500:
                            break
                        rows.append([self._json_cell(value) for value in row[:50]])
                    payload[sheet.title] = rows
            finally:
                workbook.close()
            return json.dumps(payload, ensure_ascii=False, indent=2)
        if suffix not in {".txt", ".md", ".json", ".csv", ".yaml", ".yml", ".log"}:
            raise ValueError("Supported readable files are text, Markdown, JSON, CSV, YAML, .docx, and .xlsx")
        return path.read_text(encoding="utf-8", errors="replace")[:200_000]

    @staticmethod
    def _safe_office_archive(path: Path) -> None:
        try:
            with zipfile.ZipFile(path) as archive:
                expanded = sum(item.file_size for item in archive.infolist())
                if expanded > 50 * 1024 * 1024 or len(archive.infolist()) > 10_000:
                    raise ValueError(f"{path.name} expands beyond the safe document limit")
        except zipfile.BadZipFile as error:
            raise ValueError(f"{path.name} is not a valid Office document") from error

    async def _write_file(self, arguments: dict[str, Any]) -> dict[str, Any]:
        path = self._safe_path(str(arguments["path"]))
        content = str(arguments.get("content", ""))
        return await asyncio.to_thread(self._write_file_sync, path, content)

    @staticmethod
    def _write_file_sync(path: Path, content: str) -> dict[str, Any]:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        return {"path": str(path), "bytes": len(content.encode("utf-8"))}

    async def _create_word(self, arguments: dict[str, Any]) -> dict[str, Any]:
        path = self._safe_path(str(arguments["path"]))
        if path.suffix.casefold() != ".docx":
            raise ValueError("Word documents must use the .docx extension")
        title = str(arguments.get("title", "")).strip()
        content = str(arguments.get("content", ""))
        return await asyncio.to_thread(self._create_word_sync, path, title, content)

    @staticmethod
    def _create_word_sync(path: Path, title: str, content: str) -> dict[str, Any]:
        if len(content) > 1_000_000:
            raise ValueError("Word document content is limited to one million characters per workflow step")
        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        if title:
            document.add_heading(title, level=0)
        for block in content.replace("\r\n", "\n").split("\n\n"):
            text = block.strip()
            if text:
                document.add_paragraph(text)
        document.save(path)
        return {"path": str(path), "bytes": path.stat().st_size, "format": "docx"}

    async def _create_excel(self, arguments: dict[str, Any]) -> dict[str, Any]:
        path = self._safe_path(str(arguments["path"]))
        if path.suffix.casefold() != ".xlsx":
            raise ValueError("Excel workbooks must use the .xlsx extension")
        sheet_name = str(arguments.get("sheet_name", "Data")).strip() or "Data"
        if any(character in sheet_name for character in "[]:*?/\\") or len(sheet_name) > 31:
            raise ValueError("Choose an Excel sheet name up to 31 characters without []:*?/\\")
        rows = self._spreadsheet_rows(arguments.get("data", []))
        return await asyncio.to_thread(self._create_excel_sync, path, sheet_name, rows)

    @staticmethod
    def _spreadsheet_rows(value: Any) -> list[list[Any]]:
        if isinstance(value, str):
            if len(value.encode("utf-8")) > 5 * 1024 * 1024:
                raise ValueError("Spreadsheet input is too large")
            try:
                value = json.loads(value)
            except json.JSONDecodeError:
                value = list(csv.reader(StringIO(value)))
        if isinstance(value, dict):
            value = [value]
        if not isinstance(value, list):
            value = [[value]]
        if value and all(isinstance(item, dict) for item in value):
            headers = list(dict.fromkeys(key for item in value for key in item.keys()))
            rows = [headers] + [[item.get(key, "") for key in headers] for item in value]
        else:
            rows = [item if isinstance(item, list) else [item] for item in value]
        if len(rows) > 10_000 or any(len(row) > 100 for row in rows):
            raise ValueError("A workflow can create at most 10,000 rows and 100 columns at once")
        return rows

    @staticmethod
    def _safe_cell(value: Any) -> Any:
        if value is None or isinstance(value, (bool, int, float, datetime)):
            return value
        if isinstance(value, (dict, list)):
            value = json.dumps(value, ensure_ascii=False)
        text = str(value)
        if text.startswith(("=", "+", "-", "@")):
            return "'" + text
        return text[:32_767]

    @classmethod
    def _create_excel_sync(cls, path: Path, sheet_name: str, rows: list[list[Any]]) -> dict[str, Any]:
        path.parent.mkdir(parents=True, exist_ok=True)
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = sheet_name
        for row in rows:
            sheet.append([cls._safe_cell(value) for value in row])
        if rows:
            for cell in sheet[1]:
                cell.font = Font(bold=True)
            sheet.freeze_panes = "A2"
            for index in range(1, min(len(rows[0]), 100) + 1):
                values = [str(sheet.cell(row=row, column=index).value or "") for row in range(1, min(sheet.max_row, 100) + 1)]
                sheet.column_dimensions[get_column_letter(index)].width = min(50, max(10, max(map(len, values), default=10) + 2))
        workbook.save(path)
        return {"path": str(path), "bytes": path.stat().st_size, "format": "xlsx", "rows": len(rows)}

    @staticmethod
    def _json_cell(value: Any) -> Any:
        if value is None or isinstance(value, (str, bool, int, float)):
            return value
        if isinstance(value, datetime):
            return value.isoformat()
        return str(value)

    async def _search_files(self, arguments: dict[str, Any]) -> list[dict[str, Any]]:
        query = str(arguments["query"]).casefold().strip()
        if not query:
            raise ValueError("Enter text to search for")
        return await asyncio.to_thread(self._search_files_sync, query)

    def _search_files_sync(self, query: str) -> list[dict[str, Any]]:
        matches: list[dict[str, Any]] = []
        scanned = 0
        supported = {".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".docx", ".xlsx"}
        for root in self.approved_folders:
            for path in root.rglob("*"):
                if not path.is_file() or path.suffix.casefold() not in supported:
                    continue
                scanned += 1
                if scanned > 2_000:
                    return matches
                try:
                    text = self._read_file_sync(path)
                except (OSError, ValueError):
                    continue
                index = text.casefold().find(query)
                if index >= 0:
                    matches.append({"path": str(path), "snippet": text[max(0, index - 100): index + 300]})
                if len(matches) >= 20:
                    return matches
        return matches

    async def _http_request(self, arguments: dict[str, Any]) -> dict[str, Any]:
        url = str(arguments["url"])
        if not url.startswith("https://"):
            raise ValueError("Only HTTPS endpoints are allowed")
        hostname = (urlparse(url).hostname or "").casefold()
        if not any(hostname == domain or hostname.endswith(f".{domain}") for domain in self.approved_domains):
            raise PermissionError("This domain is not in the approved HTTP allowlist")
        method = str(arguments.get("method", "GET")).upper()
        async with httpx.AsyncClient(timeout=30, follow_redirects=False) as client:
            response = await client.request(method, url, json=arguments.get("body"))
        return {"status": response.status_code, "headers": dict(response.headers), "body": response.text[:100_000]}

    @staticmethod
    def _recipients(raw: str) -> list[str]:
        recipients = [address for _, address in getaddresses([raw]) if "@" in address]
        if not recipients or len(recipients) > 10:
            raise ValueError("Enter between one and ten valid recipient email addresses")
        return recipients

    async def _send_email(self, arguments: dict[str, Any]) -> dict[str, Any]:
        return await asyncio.to_thread(self._send_email_sync, arguments)

    def _send_email_sync(self, arguments: dict[str, Any]) -> dict[str, Any]:
        config = self.email_config
        if not all(config.get(key) for key in ("host", "username", "password", "sender_email")):
            raise RuntimeError("Email is not configured. Open Settings, connect a sending account, and send a test email first.")
        recipients = self._recipients(str(arguments["to"]))
        subject = str(arguments.get("subject", "")).replace("\r", " ").replace("\n", " ")[:998]
        message = EmailMessage()
        message["From"] = formataddr((str(config.get("sender_name", "")), str(config["sender_email"])))
        message["To"] = ", ".join(recipients)
        message["Subject"] = subject
        message.set_content(str(arguments.get("body", "")))
        host = str(config["host"])
        port = int(config.get("port", 465))
        context = ssl.create_default_context()
        try:
            if config.get("security") == "starttls":
                with smtplib.SMTP(host, port, timeout=30) as smtp:
                    smtp.ehlo()
                    smtp.starttls(context=context)
                    smtp.ehlo()
                    smtp.login(str(config["username"]), str(config["password"]))
                    smtp.send_message(message)
            else:
                with smtplib.SMTP_SSL(host, port, timeout=30, context=context) as smtp:
                    smtp.login(str(config["username"]), str(config["password"]))
                    smtp.send_message(message)
        except smtplib.SMTPAuthenticationError as error:
            provider = str(config.get("provider", "custom"))
            if provider == "gmail":
                detail = (
                    "Gmail rejected the login. Use a Google App Password, not your normal Google password. "
                    "Turn on 2-Step Verification, create a 16-character App Password at "
                    "myaccount.google.com/apppasswords, save it in Settings, then send a test email."
                )
            elif provider == "yahoo":
                detail = "Yahoo rejected the login. Create and save a Yahoo app password, then send a test email."
            elif provider == "outlook":
                detail = (
                    "Microsoft rejected the login. Confirm SMTP AUTH is enabled for this account and use an app "
                    "password when required by the account administrator."
                )
            else:
                detail = "The mail server rejected the username or password. Check the SMTP credentials and send a test email."
            raise RuntimeError(detail) from error
        return {"sent": True, "to": recipients, "from": str(config["sender_email"])}

    async def _transform(self, arguments: dict[str, Any]) -> Any:
        operation = str(arguments.get("operation", "json_pretty"))
        value = arguments.get("value")
        if operation == "json_pretty":
            parsed = json.loads(value) if isinstance(value, str) else value
            return json.dumps(parsed, indent=2, ensure_ascii=False)
        if operation == "csv_to_json":
            return list(csv.DictReader(StringIO(str(value))))
        if operation == "uppercase":
            return str(value).upper()
        if operation == "lowercase":
            return str(value).lower()
        if operation == "utc_now":
            return datetime.now(UTC).isoformat()
        raise ValueError(f"Unsupported transform operation: {operation}")
